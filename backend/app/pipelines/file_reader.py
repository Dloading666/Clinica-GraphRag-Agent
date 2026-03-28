"""Document file reader - reads .docx and returns structured content"""
from dataclasses import dataclass
from typing import List
from pathlib import Path


@dataclass
class Section:
    chapter: str        # 一级标题
    section: str        # 二级标题
    content: str        # 内容文本
    document_name: str


# 标题样式名称（支持英文和中文 Word 样式）
_HEADING1_STYLES = {"heading 1", "标题 1", "标题1", "heading1"}
_HEADING2_STYLES = {"heading 2", "标题 2", "标题2", "heading2"}
_HEADING3_STYLES = {"heading 3", "标题 3", "标题3", "heading3"}


def _get_heading_level(style_name: str) -> int:
    """根据段落样式名称返回标题级别（0 表示非标题）"""
    name_lower = style_name.lower().strip()
    if name_lower in _HEADING1_STYLES:
        return 1
    if name_lower in _HEADING2_STYLES:
        return 2
    if name_lower in _HEADING3_STYLES:
        return 3
    # 通用匹配：'heading 4', '标题4' 等
    if name_lower.startswith("heading ") or name_lower.startswith("标题"):
        try:
            level = int(name_lower.replace("heading ", "").replace("标题 ", "").replace("标题", "").strip())
            return level
        except ValueError:
            pass
    return 0


def read_docx(file_path: str) -> List[Section]:
    """
    读取 Word 文档，按标题层级切分为 Section 列表。
    - 一级标题 → chapter
    - 二级标题 → section
    - 其余内容文本累积至当前 section
    """
    from docx import Document  # python-docx

    path = Path(file_path)
    doc_name = path.stem
    doc = Document(file_path)

    sections: List[Section] = []
    current_chapter = ""
    current_section = ""
    content_lines: List[str] = []

    def _flush():
        """将已累积的内容写入 sections"""
        text = "\n".join(content_lines).strip()
        if text:
            sections.append(Section(
                chapter=current_chapter,
                section=current_section,
                content=text,
                document_name=doc_name,
            ))
        content_lines.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        level = _get_heading_level(style_name)

        if level == 1:
            # 新的一级标题：先保存之前的内容
            _flush()
            current_chapter = text
            current_section = ""
        elif level == 2:
            # 新的二级标题
            _flush()
            current_section = text
        elif level >= 3:
            # 三级及以下标题作为内容处理，但加粗标注
            content_lines.append(f"【{text}】")
        else:
            # 普通正文
            content_lines.append(text)

    # 最后一段内容
    _flush()

    # 如果文档没有任何标题，把所有内容合并为一个 Section
    if not sections:
        full_text = "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        if full_text:
            sections.append(Section(
                chapter=doc_name,
                section="",
                content=full_text,
                document_name=doc_name,
            ))

    return sections


def read_document(file_path: str) -> List[Section]:
    """自动识别文件类型并读取文档"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix in [".docx", ".doc"]:
        return read_docx(file_path)
    # 未来可扩展 PDF 支持
    raise ValueError(f"Unsupported file type: {path.suffix}")
